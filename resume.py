from docx import Document
from docx.shared import Inches

def create_resume_with_photo(photo_path, resume_path):
    document = Document()

    # Insert the photo
    document.add_picture(photo_path, width=Inches(2), height=Inches(2))

    # Add the rest of the resume content
    document.add_paragraph('Your Name')
    document.add_paragraph('Contact Information')
    document.add_paragraph('Professional Summary')
    document.add_paragraph('Work Experience')
    document.add_paragraph('Education')
    document.add_paragraph('Skills')

    # Save the document
    document.save(resume_path)

# Example usage
photo_path = 'path/to/photo.jpg'
resume_path = 'path/to/resume.docx'
create_resume_with_photo(photo_path, resume_path)
